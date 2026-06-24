from docx import Document
import re
import os


def export_ieee(topic, report, references):
    doc = Document()

    # =========================
    # TITLE
    # =========================
    doc.add_heading(topic, level=1)

    # Author block
    doc.add_paragraph(
        "Pavithra Shine\n"
        "Department of Computer Engineering\n"
        "XYZ College\n"
        "Email: example@gmail.com"
    )

    # =========================
    # SECTION MAPPING
    # =========================
    sections = [
        ("1. Executive Summary", "Abstract"),
        ("2. Introduction", "Introduction"),
        ("3. Main Findings", "Results and Discussion"),
        ("4. Challenges and Limitations", "Limitations"),
        ("5. Future Outlook", "Future Work"),
        ("6. Conclusion", "Conclusion")
    ]

    # =========================
    # EXTRACT AND WRITE SECTIONS
    # =========================
    for i, (old_name, ieee_name) in enumerate(sections):

        if old_name not in report:
            continue

        # Add IEEE heading
        doc.add_heading(ieee_name, level=2)

        # Start position (skip heading text itself)
        start = report.find(old_name) + len(old_name)

        # End position = next section start or end of report
        end = len(report)

        for next_old, _ in sections[i + 1:]:
            pos = report.find(next_old, start)
            if pos != -1:
                end = min(end, pos)

        # Extract content
        content = report[start:end].strip()

        # Clean markdown artifacts
        content = re.sub(r"#+", "", content)
        content = content.replace("---", "")
        content = content.strip()

        doc.add_paragraph(content)

    # =========================
    # REFERENCES SECTION
    # =========================
    doc.add_heading("References", level=2)

    for i, ref in enumerate(references, start=1):
        doc.add_paragraph(f"[{i}] {ref}")

    # =========================
    # SAFE FILE NAME GENERATION
    # =========================
    safe_topic = "".join(
        c if c.isalnum() or c in (" ", "_", "-") else "_"
        for c in topic
    ).replace(" ", "_")

    filename = f"{safe_topic}.docx"

    counter = 1
    while os.path.exists(filename):
        filename = f"{safe_topic}({counter}).docx"
        counter += 1

    # =========================
    # SAVE FILE
    # =========================
    doc.save(filename)
    print(f"IEEE report saved as {filename}")